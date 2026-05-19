# reports/bk_docx.py
# ====================================================================
# Template de documento Word padrao BK Engenharia e Tecnologia
# Gera relatorios .docx com header, footer, estilos e equacoes OMML
# ====================================================================

from __future__ import annotations

import os
import io
import math
from typing import Optional, List, Tuple
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

# Cores BK
BK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
BK_DARK = RGBColor(0x33, 0x33, 0x33)
BK_GRAY = RGBColor(0x66, 0x66, 0x66)
BK_LIGHT = RGBColor(0xD5, 0xE8, 0xF0)
BK_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

ASSETS_DIR = Path(__file__).parent.parent / "assets"


# ====================================================================
# Classe principal
# ====================================================================

class BKReport:
    """Gerador de relatorio Word padrao BK Engenharia e Tecnologia."""

    def __init__(
        self,
        titulo_estudo: str,
        codigo_doc: str = "",
        revisao: str = "0A",
        aprovacao: str = "Eng. Eletricista",
    ):
        self.doc = Document()
        self.titulo_estudo = titulo_estudo
        self.codigo_doc = codigo_doc
        self.revisao = revisao
        self.aprovacao = aprovacao
        self._figure_count = 0
        self._table_count = 0
        self._eq_count = 0

        self._setup_page()
        self._setup_styles()
        self._setup_header()
        self._setup_footer()

    # ----------------------------------------------------------------
    # Configuracao de pagina (A4, margens conforme ref)
    # ----------------------------------------------------------------
    def _setup_page(self):
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)

    # ----------------------------------------------------------------
    # Estilos (Body: Arial 12pt, Heading1: Arial 12pt bold caps, etc.)
    # ----------------------------------------------------------------
    def _setup_styles(self):
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = BK_DARK
        pf = style.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.15
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Heading 1 — conforme modelo: 16pt bold azul, espaço antes 24pt, depois 0
        h1 = self.doc.styles["Heading 1"]
        h1.font.name = "Calibri"
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.all_caps = False
        h1.font.color.rgb = BK_BLUE
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after = Pt(0)
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Heading 2 — conforme modelo: 13pt bold azul, espaço antes 10pt, depois 0
        h2 = self.doc.styles["Heading 2"]
        h2.font.name = "Calibri"
        h2.font.size = Pt(13)
        h2.font.bold = True
        h2.font.underline = False
        h2.font.color.rgb = BK_BLUE
        h2.paragraph_format.space_before = Pt(10)
        h2.paragraph_format.space_after = Pt(0)
        h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Heading 3 — conforme modelo: 11pt bold azul, espaço antes 10pt, depois 0
        h3 = self.doc.styles["Heading 3"]
        h3.font.name = "Calibri"
        h3.font.size = Pt(11)
        h3.font.bold = True
        h3.font.italic = False
        h3.font.color.rgb = BK_BLUE
        h3.paragraph_format.space_before = Pt(10)
        h3.paragraph_format.space_after = Pt(0)

    # ----------------------------------------------------------------
    # Header — REMOVIDO conforme modelo de referência (sem cabeçalho)
    # ----------------------------------------------------------------
    def _setup_header(self):
        section = self.doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        # Limpar qualquer parágrafo existente para garantir cabeçalho vazio
        for p in header.paragraphs:
            p.clear()

    # ----------------------------------------------------------------
    # Footer com numero de pagina — conforme modelo (direita, 9pt)
    # ----------------------------------------------------------------
    def _setup_footer(self):
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_pre = p.add_run("Página ")
        run_pre.font.size = Pt(8)
        run_pre.font.name = "Calibri"
        run_pre.font.color.rgb = BK_GRAY
        # campo PAGE
        run_pg = p.add_run()
        run_pg.font.size = Pt(8)
        fld_char1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_pg._r.append(fld_char1)
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run_pg._r.append(instr)
        fld_char2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_pg._r.append(fld_char2)

    # ----------------------------------------------------------------
    # Metodos auxiliares de formatacao
    # ----------------------------------------------------------------
    def _set_cell_border(self, cell, sz=4, color="A0A0A0"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)

    def _shade_cell(self, cell, color="D5E8F0"):
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)

    # ================================================================
    # API publica — conteudo
    # ================================================================

    def add_heading1(self, text: str):
        """Adiciona titulo principal (Heading 1)."""
        self.doc.add_heading(text, level=1)

    def add_heading2(self, text: str):
        """Adiciona subtitulo (Heading 2)."""
        self.doc.add_heading(text, level=2)

    def add_heading3(self, text: str):
        """Adiciona sub-subtitulo (Heading 3)."""
        self.doc.add_heading(text, level=3)

    def add_body(self, text: str, bold: bool = False, italic: bool = False):
        """Adiciona paragrafo de corpo de texto."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.space_before = Pt(4)
        pf.space_after = Pt(4)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.bold = bold
        run.font.italic = italic
        return p

    def add_body_list(self, items: list[str], numbered: bool = False):
        """Adiciona lista de itens."""
        for i, item in enumerate(items):
            prefix = f"{i+1}. " if numbered else "\u2022 "
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"{prefix}{item}")
            run.font.name = "Arial"
            run.font.size = Pt(11)

    def add_equation(self, omml_element, caption: str = ""):
        """Insere equacao OMML centrada no documento."""
        self._eq_count += 1
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        # Injeta o OMML no paragrafo
        p._p.append(omml_element)

        if caption:
            pc = self.doc.add_paragraph()
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rc = pc.add_run(f"({self._eq_count})  {caption}")
            rc.font.size = Pt(9)
            rc.font.italic = True
            rc.font.name = "Arial"
            rc.font.color.rgb = BK_GRAY

    def add_figure_from_buf(self, buf: io.BytesIO, caption: str = "", width_cm: float = 14.0):
        """Insere figura a partir de buffer de imagem PNG."""
        self._figure_count += 1
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        buf.seek(0)
        run.add_picture(buf, width=Cm(width_cm))

        if caption:
            pc = self.doc.add_paragraph()
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pc.paragraph_format.space_after = Pt(8)
            rc = pc.add_run(f"Figura {self._figure_count} – {caption}")
            rc.font.size = Pt(10)
            rc.font.italic = True
            rc.font.name = "Arial"
            rc.font.color.rgb = BK_GRAY

    def add_figure_from_plotly(self, fig, caption: str = "", width_cm: float = 14.0):
        """Insere figura Plotly como imagem PNG."""
        try:
            buf = io.BytesIO()
            fig.write_image(buf, format="png", width=900, height=500, scale=2)
            self.add_figure_from_buf(buf, caption, width_cm)
        except Exception:
            self.add_body(f"[Figura: {caption} - exportação Plotly indisponível]", italic=True)

    def add_figure_from_matplotlib(self, fig, caption: str = "", width_cm: float = 14.0):
        """Insere figura matplotlib como imagem PNG."""
        buf = io.BytesIO()
        fig.tight_layout()
        fig.savefig(buf, format="png", dpi=200, bbox_inches="tight")
        plt.close(fig)
        self.add_figure_from_buf(buf, caption, width_cm)

    def add_result_table(self, headers: list[str], rows: list[list[str]], caption: str = ""):
        """Adiciona tabela de resultados formatada."""
        self._table_count += 1

        if caption:
            pc = self.doc.add_paragraph()
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pc.paragraph_format.space_before = Pt(8)
            rc = pc.add_run(f"Tabela {self._table_count} – {caption}")
            rc.font.size = Pt(10)
            rc.font.bold = True
            rc.font.name = "Arial"
            rc.font.color.rgb = BK_BLUE

        ncols = len(headers)
        nrows = len(rows)
        tbl = self.doc.add_table(rows=nrows + 1, cols=ncols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for j, h in enumerate(headers):
            cell = tbl.cell(0, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = "Arial"
            run.font.color.rgb = BK_WHITE
            self._shade_cell(cell, "084C89")
            self._set_cell_border(cell, sz=4, color="084C89")

        # Data rows
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = tbl.cell(i + 1, j)
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(val))
                run.font.size = Pt(9)
                run.font.name = "Arial"
                self._set_cell_border(cell)
                if i % 2 == 0:
                    self._shade_cell(cell, "F0F5FA")

    def add_kpi_table(self, kpis: list[tuple[str, str, str]]):
        """Adiciona tabela de KPIs: [(label, valor, unidade), ...]."""
        ncols = min(len(kpis), 4)
        nrows = math.ceil(len(kpis) / ncols)
        tbl = self.doc.add_table(rows=nrows * 2, cols=ncols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        for idx, (label, valor, unidade) in enumerate(kpis):
            r = (idx // ncols) * 2
            c = idx % ncols
            # Label
            cell_l = tbl.cell(r, c)
            cell_l.text = ""
            p = cell_l.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(label)
            run.font.size = Pt(8)
            run.font.name = "Arial"
            run.font.bold = True
            run.font.color.rgb = BK_BLUE
            self._shade_cell(cell_l, "E8F0FE")
            self._set_cell_border(cell_l, sz=2, color="B0C4DE")
            # Valor
            cell_v = tbl.cell(r + 1, c)
            cell_v.text = ""
            p = cell_v.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{valor} {unidade}")
            run.font.size = Pt(11)
            run.font.name = "Arial"
            run.font.bold = True
            self._set_cell_border(cell_v, sz=2, color="B0C4DE")

    def add_page_break(self):
        """Adiciona quebra de pagina."""
        self.doc.add_page_break()

    def add_toc(self):
        """Adiciona sumario (Table of Contents). Precisa atualizar no Word."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        fld_char1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fld_char1)
        instr = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve">'
            ' TOC \\o "1-3" \\h \\z \\u </w:instrText>'
        )
        run._r.append(instr)
        fld_char2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run._r.append(fld_char2)
        run2 = p.add_run("(Atualizar sumário: clique com botão direito → Atualizar campo)")
        run2.font.size = Pt(10)
        run2.font.color.rgb = BK_GRAY
        fld_char3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run2._r.append(fld_char3)

        self.add_body(
            "Nota: Ao abrir o documento no Microsoft Word, clique com o botão direito sobre o "
            "sumário acima e selecione \"Atualizar Campo\" → \"Atualizar índice inteiro\" "
            "para gerar a numeração de páginas.",
            italic=True
        )

    # ================================================================
    # Salvar
    # ================================================================
    def save(self, path: str):
        """Salva o documento .docx."""
        self.doc.save(path)

    def to_bytes(self) -> bytes:
        """Retorna o documento como bytes (para download no Streamlit)."""
        buf = io.BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    # ================================================================
    # Metodos de conveniencia para secoes padrao de relatorio
    # ================================================================
    def add_cover(self, titulo: str, subtitulo: str = ""):
        """Pagina de capa."""
        for _ in range(6):
            self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(titulo.upper())
        run.font.name = "Arial"
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = BK_BLUE

        if subtitulo:
            p2 = self.doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(subtitulo)
            run2.font.name = "Arial"
            run2.font.size = Pt(14)
            run2.font.color.rgb = BK_GRAY

        for _ in range(4):
            self.doc.add_paragraph()

        p3 = self.doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run("Software BK Estudos Elétricos")
        run3.font.name = "Arial"
        run3.font.size = Pt(11)
        run3.font.italic = True

        p4 = self.doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run4 = p4.add_run("BK Engenharia e Tecnologia")
        run4.font.name = "Arial"
        run4.font.size = Pt(11)
        run4.font.bold = True
        run4.font.color.rgb = BK_BLUE

    def add_references_section(self, refs: list[str]):
        """Secao de referencias bibliograficas."""
        self.add_heading1("Referências Bibliográficas")
        for i, ref in enumerate(refs, 1):
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-1.0)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"[{i}] {ref}")
            run.font.name = "Arial"
            run.font.size = Pt(10)

    # ── Seção paisagem (landscape) para tabelas largas ──────────
    def start_landscape_section(self):
        """Inicia nova seção em orientação paisagem (landscape)."""
        new_section = self.doc.add_section(2)  # 2 = NEW_PAGE
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = Cm(29.7)
        new_section.page_height = Cm(21.0)
        new_section.top_margin = Cm(2.0)
        new_section.bottom_margin = Cm(1.5)
        new_section.left_margin = Cm(2.0)
        new_section.right_margin = Cm(2.0)

    def end_landscape_section(self):
        """Retorna orientação retrato (portrait)."""
        new_section = self.doc.add_section(2)
        new_section.orientation = WD_ORIENT.PORTRAIT
        new_section.page_width = Cm(21.0)
        new_section.page_height = Cm(29.7)
        new_section.top_margin = Cm(3.5)
        new_section.bottom_margin = Cm(2.0)
        new_section.left_margin = Cm(3.0)
        new_section.right_margin = Cm(2.0)

    def add_compact_table(self, headers: list[str], rows: list[list[str]],
                          caption: str = "", font_size: int = 7):
        """Tabela compacta com fonte pequena para tabelas grandes (ex: flecha × vão × temp)."""
        self._table_count += 1

        if caption:
            pc = self.doc.add_paragraph()
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pc.paragraph_format.space_before = Pt(6)
            rc = pc.add_run(f"Tabela {self._table_count} – {caption}")
            rc.font.size = Pt(9)
            rc.font.bold = True
            rc.font.name = "Arial"
            rc.font.color.rgb = BK_BLUE

        ncols = len(headers)
        nrows = len(rows)
        tbl = self.doc.add_table(rows=nrows + 1, cols=ncols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = True

        # Header
        for j, h in enumerate(headers):
            cell = tbl.cell(0, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(h)
            run.font.bold = True
            run.font.size = Pt(font_size)
            run.font.name = "Arial"
            run.font.color.rgb = BK_WHITE
            self._shade_cell(cell, "084C89")
            self._set_cell_border(cell, sz=2, color="084C89")

        # Data
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = tbl.cell(i + 1, j)
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(str(val))
                run.font.size = Pt(font_size)
                run.font.name = "Arial"
                self._set_cell_border(cell, sz=2)
                if j == 0:
                    run.font.bold = True
                    self._shade_cell(cell, "E8EFF5")
                elif i % 2 == 0:
                    self._shade_cell(cell, "F7F9FC")

    def add_figure_from_base64(self, b64_str: str, caption: str = "", width_cm: float = 14.0):
        """Insere imagem a partir de string base64."""
        import io, base64
        img_data = base64.b64decode(b64_str)
        buf = io.BytesIO(img_data)
        self.add_figure_from_buf(buf, caption, width_cm)
